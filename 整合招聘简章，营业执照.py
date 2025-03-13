import os
import re
import os
from io import BytesIO

import requests
from bs4 import BeautifulSoup
from tabulate import tabulate
from docx import Document
from docx.shared import Inches
from urllib.parse import urljoin
import requests
from bs4 import BeautifulSoup
from docx import Document
from lxml import etree # 用于将相对路径转换为绝对路径
def parse_table(table):
    """解析表格并格式化为整齐的文本"""
    headers = []
    rows = []
    for row in table.find_all("tr"):
        cols = row.find_all(["td", "th"])
        cols_text = [col.get_text(strip=True) for col in cols]
        if not headers:
            headers = cols_text  # 获取表头
        else:
            rows.append(cols_text)  # 获取数据行

    # 用 tabulate 格式化表格
    table_str = tabulate(rows, headers=headers, tablefmt="grid", stralign="center")
    return table_str


def download_image(img_url, save_dir="images"):
    """下载图片并保存到本地目录"""
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)

    try:
        # 将相对路径转为绝对路径
        img_url = urljoin("http://example.com", img_url)  # 这里替换成实际的基础 URL

        img_data = requests.get(img_url).content
        img_name = os.path.join(save_dir, img_url.split("/")[-1])

        # 确保图片路径有效
        with open(img_name, "wb") as handler:
            handler.write(img_data)

        # 返回图片的绝对路径
        return os.path.abspath(img_name)
    except Exception as e:
        print(f"下载图片失败: {e}")
        return None
base_folder = r"F:\pycharm项目\测试\就业办网站脚本test\3.26双选"
for i in range(1,28):
    with open(f"{i}.html", "r", encoding="utf-8") as f:
        html_content = f.read()
    tree = etree.HTML(html_content)
    name_elements = tree.xpath('//*[@id="app"]/div[1]/div[1]/div/div/div[2]/div/p')
    if name_elements:
        name = name_elements[0].text.strip()  # 提取第一个 <p> 标签的文本并去除空白字符
    else:
        name = "未知公司"
    photo_img = tree.xpath('//*[@id="pane-base"]/div[10]/div/div[2]/div/div/img')
    if photo_img:
        img_url = photo_img[0].get("src")
        print(f"公司名称: {name}")
        print(f"特定图片 URL: {img_url}")

        # 创建公司文件夹
        company_folder = os.path.join(base_folder, name)
        os.makedirs(company_folder, exist_ok=True)

        # 下载图片
        try:
            response = requests.get(img_url)
            response.raise_for_status()  # 检查请求是否成功

            # 生成图片保存路径
            img_extension = os.path.splitext(img_url)[1]  # 获取文件扩展名（如 .jpg, .png）
            img_name = f"{name}-营业执照{img_extension}"  # 新文件名：公司名-营业执照.扩展名
            save_path = os.path.join(company_folder, img_name)

            # 保存图片
            with open(save_path, "wb") as f:
                f.write(response.content)

            print(f"图片已保存到: {save_path}")
        except Exception as e:
            print(f"下载失败: {e}")
    else:
        print(f"未找到符合条件的图片（文件: {i}.html）")

    with open(f"招聘简章{i}.html", "r", encoding="utf-8") as file:
        html_content = file.read()

    # 解析 HTML
    soup = BeautifulSoup(html_content, "html.parser")

    # 定位目标 div
    target_div = soup.select_one(
        "#pane-apply > div > div:nth-of-type(1) > div:nth-of-type(12) > div > div:nth-of-type(2) > div")
    if not target_div:
        target_div = soup.select_one(
            "#pane-apply > div > div:nth-of-type(1) > div:nth-of-type(11) > div > div:nth-of-type(2) > div")

    if target_div or target_div :
        content_list = []

        # 创建 Word 文档
        doc = Document()

        for elem in target_div.contents:
            if elem.name == "table":
                parsed_table = parse_table(elem)
                content_list.append("\n[表格]\n" + parsed_table)  # 格式化表格
            elif elem.name == "img":
                img_url = elem.get("src")
                img_path = download_image(img_url, company_folder)  # 下载并保存图片
                if img_path:
                    content_list.append(f"\n[图片: {img_path}]")
                    image_bytes = BytesIO(requests.get(img_url).content)
                    doc.add_picture(image_bytes, width=Inches(4.0))
            elif elem.name:
                content_list.append("\n" + elem.get_text(strip=True))
            else:
                content_list.append(str(elem))

        # 合并文本内容
        result = "\n".join(content_list)
        doc.add_paragraph(result)

        # 保存到 Word 文件
        word_file_path = os.path.join(company_folder, f"{name}-招聘简章.docx")
        doc.save(word_file_path)
        print(f"已成功保存到 Word 文件: {word_file_path}")
        print(f"第{i}个word成功！")

    else:
        print(f"未找到目标标签（文件: 招聘简章{i}.html）")
        print(f"第{i}个word失败！")
