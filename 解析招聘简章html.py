import os
from io import BytesIO

import requests
from bs4 import BeautifulSoup
from tabulate import tabulate
from docx import Document
from docx.shared import Inches
from urllib.parse import urljoin  # 用于将相对路径转换为绝对路径

# 读取 HTML 文件
with open(f"招聘简章{i}.html", "r", encoding="utf-8") as file:
    html_content = file.read()

# 解析 HTML
soup = BeautifulSoup(html_content, "html.parser")

# 定位目标 div
target_div = soup.select_one(
    "#pane-apply > div > div:nth-of-type(1) > div:nth-of-type(12) > div > div:nth-of-type(2) > div")


def parse_table(table):
    """解析表格并格式化为整齐的文本"""
    headers = []
    rows = []
    for row in table.find_all("tr"):
        cols = row.find_all(["td", "th"])
        cols_text = [col.get_text(strip=True) for col in cols]
        if not headers:
            headers = cols_text  # 获取表头
        else:
            rows.append(cols_text)  # 获取数据行

    # 用 tabulate 格式化表格
    table_str = tabulate(rows, headers=headers, tablefmt="grid", stralign="center")
    return table_str


def download_image(img_url, save_dir="images"):
    """下载图片并保存到本地目录"""
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)

    try:
        # 将相对路径转为绝对路径
        img_url = urljoin("http://example.com", img_url)  # 这里替换成实际的基础 URL

        img_data = requests.get(img_url).content
        img_name = os.path.join(save_dir, img_url.split("/")[-1])

        # 确保图片路径有效
        with open(img_name, "wb") as handler:
            handler.write(img_data)

        # 返回图片的绝对路径
        return os.path.abspath(img_name)
    except Exception as e:
        print(f"下载图片失败: {e}")
        return None


# 创建 Word 文档
doc = Document()

if target_div:
    content_list = []

    for elem in target_div.contents:
        if elem.name == "table":
            parsed_table = parse_table(elem)
            content_list.append("\n[表格]\n" + parsed_table)  # 格式化表格
        elif elem.name == "img":
            img_url = elem.get("src")
            response = requests.get(img_url)
            image_bytes = BytesIO(response.content)
            doc.add_picture(image_bytes, width=Inches(4.0))
            # if img_url:
            #     img_path = download_image(img_url)
            #     if img_path:
            #         content_list.append("\n[图片]")
            #         doc.add_paragraph(f"插入图片：{img_url}")
            #         # 确保插入的是图片文件的路径
            #         doc.add_picture(image_bytes, width=Inches(4.0))   # 插入图片并调整大小
        elif elem.name:
            content_list.append("\n" + elem.get_text(strip=True))
        else:
            content_list.append(str(elem))

            # 合并文本内容
    result = "\n".join(content_list)
    doc.add_paragraph(result)

    # 保存到 Word 文件
    doc.save(f"{name}-招聘简章.docx")
    print("已成功保存到 Word 文件")
else:
    print("未找到目标标签")







