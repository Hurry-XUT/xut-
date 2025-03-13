import os
import re
from bs4 import BeautifulSoup
from docx import Document
from lxml import etree
main_folder = "3.20双选"
os.makedirs(main_folder, exist_ok=True)
for i in range(1, 4):  #三页html,(调整)
    with open(f"page_content_{i}.html", "r", encoding="utf-8") as f:
        html_content = f.read()
    tree = etree.HTML(html_content)
    # 提取原始文本列表
    raw_items = tree.xpath('//*[@id="companyListStyle"]/div[1]/div[3]/table/tbody/tr[1]/td[4]/div/div[2]/span/div/text()')
    for i in range(1, 11):  # XPath 索引从 1 开始
        # 动态生成 XPath（替换 tr[1] 为 tr[i]）
        raw_items = tree.xpath(f'//*[@id="companyListStyle"]/div[1]/div[3]/table/tbody/tr[{i}]/td[4]/div/div[2]/span/div/text()')
        cleaned_items = [item.strip() for item in raw_items if item.strip() != ""]
        print(cleaned_items)
        for company in cleaned_items:
            company_folder = os.path.join(main_folder, company)
            os.makedirs(company_folder, exist_ok=True)  # 创建企业子文件夹
            # 2. 创建 Word 文档
            word_file_path = os.path.join(company_folder, f"{company}-招聘简章.docx")
            doc = Document()
            # doc.add_heading(f"{company} 招聘简章", level=1)
            # doc.add_paragraph("这里填写公司的招聘信息，可以后续补充详细内容。")
            doc.save(word_file_path)  # 保存 Word 文件
            print(f"创建 Word 文档：{word_file_path}")

